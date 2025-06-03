from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches
import time
import requests
import os

driver = webdriver.Chrome()
driver.get("https://www.dpm.org.cn/collection/foreigns.html")

datalist = driver.find_element(By.ID, "datalist")
li_divs = datalist.find_elements(By.CLASS_NAME, "li")

doc = Document()

for idx, li in enumerate(li_divs):
    class_attr = li.get_attribute("class")
    if "li-th" in class_attr:
        continue
    try:
        a_tag = li.find_element(By.TAG_NAME, "a")
        href = a_tag.get_attribute("href")
        if href:
            if href.startswith("/"):
                href = "https://www.dpm.org.cn" + href
            driver.execute_script("window.open('{}');".format(href))
            driver.switch_to.window(driver.window_handles[-1])
            time.sleep(1)

            # 获取title
            title = driver.title

            # 获取meta[name="Description"]的content
            try:
                meta = driver.find_element(By.XPATH, '//meta[@name="Description"]')
                content = meta.get_attribute("content")
            except:
                content = ""

            # 获取图片src
            try:
                img = driver.find_element(By.XPATH, '//img[contains(@src, "/Uploads/Picture/")]')
                img_url = img.get_attribute("src")
                # 下载图片
                img_data = requests.get(img_url).content
                img_name = f"img_{idx}.jpg"
                with open(img_name, "wb") as f:
                    f.write(img_data)
                # 写入Word
                doc.add_heading(title, level=2)
                doc.add_paragraph(content)
                doc.add_picture(img_name, width=Inches(3))
                os.remove(img_name)
            except:
                # 没有图片也写入标题和内容
                doc.add_heading(title, level=2)
                doc.add_paragraph(content)

            driver.close()
            driver.switch_to.window(driver.window_handles[0])
    except:
        continue

driver.quit()
doc.save("故宫外国文物详情.docx")